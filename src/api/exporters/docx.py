"""DOCX exporter — walks the markdown-it AST and emits python-docx elements.

Handles the report shapes the synthesizer emits: headings (h1-h3),
paragraphs with inline strong/em/code/links, bullet + ordered
lists, block quotes, code blocks, and GFM tables.
"""

from __future__ import annotations

import io
from datetime import UTC, datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token

from src.api.jobs import Job


def render_docx(job: Job) -> bytes:
    """Return the job's report as DOCX bytes."""
    doc = Document()
    _apply_default_style(doc)

    _write_title(doc, job)
    _write_metadata_table(doc, job)

    body = job.result or "(no report body)"
    tokens = MarkdownIt("commonmark", {"html": False}).enable("table").parse(body)
    _write_tokens(doc, tokens)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Styles + top matter
# ---------------------------------------------------------------------------


def _apply_default_style(doc: Any) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _write_title(doc: Any, job: Job) -> None:
    title = doc.add_paragraph()
    title_run = title.add_run("Research briefing")
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(job.query)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def _write_metadata_table(doc: Any, job: Job) -> None:
    rows: list[tuple[str, str]] = [("Job ID", job.job_id)]
    if job.completed_at is not None:
        rows.append(
            (
                "Completed",
                datetime.fromtimestamp(job.completed_at, tz=UTC).isoformat(
                    timespec="seconds"
                ),
            )
        )
    if job.iterations is not None:
        rows.append(("Iterations", str(job.iterations)))
    if job.quality_score is not None:
        rows.append(("Quality", f"{job.quality_score:.2f}"))
    if job.cost_usd is not None:
        rows.append(("Cost", f"${job.cost_usd:.4f}"))
    elapsed = job.elapsed_sec()
    if elapsed is not None:
        rows.append(("Elapsed", f"{elapsed:.1f}s"))

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(rows):
        row_cells = table.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = value
        # Label column: bold + small caps look.
        for para in row_cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)

    doc.add_paragraph()  # blank line after metadata


# ---------------------------------------------------------------------------
# Token walker
# ---------------------------------------------------------------------------


def _write_tokens(doc: Any, tokens: list[Token]) -> None:
    i = 0
    while i < len(tokens):
        tok = tokens[i]

        if tok.type == "heading_open":
            level = int(tok.tag[1:])
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            heading = doc.add_heading(level=min(level, 4))
            heading.paragraph_format.space_before = Pt(10)
            _apply_inline(heading, inline)
            i += 3
            continue

        if tok.type == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            para = doc.add_paragraph()
            _apply_inline(para, inline)
            i += 3
            continue

        if tok.type in ("bullet_list_open", "ordered_list_open"):
            style = (
                "List Bullet"
                if tok.type == "bullet_list_open"
                else "List Number"
            )
            consumed = _write_list(doc, tokens, i, style)
            i += consumed
            continue

        if tok.type == "blockquote_open":
            inner: list[Token] = []
            depth = 1
            j = i + 1
            while j < len(tokens) and depth > 0:
                if tokens[j].type == "blockquote_open":
                    depth += 1
                elif tokens[j].type == "blockquote_close":
                    depth -= 1
                    if depth == 0:
                        break
                inner.append(tokens[j])
                j += 1
            for k, t in enumerate(inner):
                if t.type == "paragraph_open":
                    inline = inner[k + 1] if k + 1 < len(inner) else None
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Pt(24)
                    _apply_inline(para, inline)
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            i = j + 1
            continue

        if tok.type in ("fence", "code_block"):
            code = tok.content.rstrip("\n")
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(24)
            run = para.add_run(code)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            i += 1
            continue

        if tok.type == "hr":
            doc.add_paragraph("____________________").alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )
            i += 1
            continue

        if tok.type == "table_open":
            consumed = _write_table(doc, tokens, i)
            i += consumed
            continue

        i += 1


def _write_list(doc: Any, tokens: list[Token], start: int, style: str) -> int:
    j = start + 1
    depth = 1
    while j < len(tokens) and depth > 0:
        tok = tokens[j]
        if tok.type in ("bullet_list_open", "ordered_list_open"):
            depth += 1
            j += 1
            continue
        if tok.type in ("bullet_list_close", "ordered_list_close"):
            depth -= 1
            j += 1
            if depth == 0:
                break
            continue

        if tok.type == "list_item_open" and depth == 1:
            k = j + 1
            para = doc.add_paragraph(style=style)
            item_depth = 1
            first_para = True
            while k < len(tokens) and item_depth > 0:
                if tokens[k].type == "list_item_open":
                    item_depth += 1
                elif tokens[k].type == "list_item_close":
                    item_depth -= 1
                    if item_depth == 0:
                        break
                elif tokens[k].type == "paragraph_open":
                    inline = tokens[k + 1] if k + 1 < len(tokens) else None
                    if first_para:
                        _apply_inline(para, inline)
                        first_para = False
                    else:
                        sub = doc.add_paragraph()
                        _apply_inline(sub, inline)
                    k += 2
                k += 1
            j = k + 1
            continue
        j += 1
    return j - start


def _write_table(doc: Any, tokens: list[Token], start: int) -> int:
    rows: list[list[str]] = []
    row_is_header: list[bool] = []
    current: list[str] = []
    current_is_header = False
    j = start + 1
    depth = 1
    in_row = False
    while j < len(tokens) and depth > 0:
        tok = tokens[j]
        if tok.type == "table_close":
            depth -= 1
            j += 1
            continue
        if tok.type == "tr_open":
            in_row = True
            current = []
            current_is_header = False
        elif tok.type == "tr_close":
            in_row = False
            rows.append(current)
            row_is_header.append(current_is_header)
        elif tok.type in ("th_open", "td_open") and in_row:
            if tok.type == "th_open":
                current_is_header = True
            inline = tokens[j + 1] if j + 1 < len(tokens) else None
            current.append(_inline_to_plain(inline))
        j += 1

    if not rows:
        return j - start

    n_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < n_cols:
            r.append("")

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for row_idx, row in enumerate(rows):
        cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(row):
            cells[col_idx].text = text
            if row_is_header[row_idx]:
                for para in cells[col_idx].paragraphs:
                    for run in para.runs:
                        run.bold = True
                # Header row background — a light gray. python-docx
                # exposes this via low-level oxml manipulation.
                tc_pr = cells[col_idx]._tc.get_or_add_tcPr()
                shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): "E2E8F0"})
                tc_pr.append(shd)
    doc.add_paragraph()
    return j - start


# ---------------------------------------------------------------------------
# Inline rendering
# ---------------------------------------------------------------------------


def _apply_inline(para: Any, inline: Token | None) -> None:
    """Render an inline token's children as `docx.Run` objects on `para`."""
    if inline is None or inline.children is None:
        return
    # Values are `bool` for the boolean style flags (bold, italic,
    # code) and `str` for the link href. `Any` keeps the stack
    # heterogeneous without a Union tax at every call site.
    style_stack: list[dict[str, Any]] = [{}]
    for child in inline.children:
        t = child.type
        style = style_stack[-1]
        if t == "text":
            _add_run(para, child.content, style)
        elif t == "code_inline":
            _add_run(para, child.content, {**style, "code": True})
        elif t == "strong_open":
            style_stack.append({**style, "bold": True})
        elif t == "strong_close":
            if len(style_stack) > 1:
                style_stack.pop()
        elif t == "em_open":
            style_stack.append({**style, "italic": True})
        elif t == "em_close":
            if len(style_stack) > 1:
                style_stack.pop()
        elif t == "link_open":
            href = child.attrGet("href") or ""
            style_stack.append({**style, "link": href})
        elif t == "link_close":
            if len(style_stack) > 1:
                style_stack.pop()
        elif t in ("softbreak", "hardbreak"):
            para.add_run().add_break()


def _add_run(para: Any, text: str, style: dict[str, Any]) -> None:
    if not text:
        return
    run = para.add_run(text)
    if style.get("bold"):
        run.bold = True
    if style.get("italic"):
        run.italic = True
    if style.get("code"):
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    if style.get("link"):
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        run.underline = True


def _inline_to_plain(token: Token | None) -> str:
    """Flatten an inline token to plain text — used for table cells
    where python-docx's Table.cell.text is single-run only."""
    if token is None or token.children is None:
        return ""
    return "".join(c.content for c in token.children if c.content)
